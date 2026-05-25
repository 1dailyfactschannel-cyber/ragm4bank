import uuid
import time
import os
from typing import Optional
from sqlalchemy.ext.asyncio import AsyncSession
from app.db.models.models import Document, DocumentChunk
from app.rag.chunking.strategies import ChunkingStrategy
from app.rag.embeddings.ollama_embeddings import OllamaEmbeddings
from app.utils.logging import setup_logger

logger = setup_logger()


class DocumentProcessor:
    def __init__(self):
        self.chunking_strategy = ChunkingStrategy()
        self.embeddings = OllamaEmbeddings()
    
    async def process_document(
        self, 
        db: AsyncSession, 
        file_path: str, 
        title: str,
        uploaded_by: Optional[str] = None
    ) -> dict:
        """
        Обрабатывает документ: извлекает текст, чанкует, векторизует и сохраняет в БД
        """
        try:
            # Определяем тип файла
            file_ext = file_path.split('.')[-1].lower()
            
            # Извлекаем текст
            text = await self._extract_text(file_path, file_ext)
            
            # Чанкуем
            chunks_with_metadata = self.chunking_strategy.chunk_with_metadata(text, file_ext)
            
            # Создаем документ в БД
            document_id = uuid.uuid4()
            document = Document(
                id=document_id,
                title=title,
                filename=os.path.basename(file_path),
                file_path=file_path,
                uploaded_by=uploaded_by
            )
            db.add(document)
            await db.flush()
            
            # Создаем эмбеддинги для чанков
            chunk_texts = [chunk["chunk_text"] for chunk in chunks_with_metadata]
            embeddings = await self.embeddings.embed_batch(chunk_texts)
            
            # Сохраняем чанки с эмбеддингами
            for i, (chunk_data, embedding) in enumerate(zip(chunks_with_metadata, embeddings)):
                doc_chunk = DocumentChunk(
                    id=uuid.uuid4(),
                    document_id=document_id,
                    chunk_text=chunk_data["chunk_text"],
                    chunk_index=chunk_data["chunk_index"],
                    embedding=embedding,
                    chunk_metadata=chunk_data["metadata"]
                )
                db.add(doc_chunk)
            
            await db.commit()
            
            logger.info(f"Processed document '{title}' into {len(chunks_with_metadata)} chunks")
            
            return {
                "document_id": document_id,
                "chunks_count": len(chunks_with_metadata),
                "status": "success"
            }
            
        except Exception as e:
            await db.rollback()
            logger.error(f"Error processing document: {str(e)}")
            raise
    
    async def _extract_text(self, file_path: str, file_ext: str) -> str:
        """
        Извлекает текст из файла в зависимости от расширения
        """
        try:
            if file_ext == "pdf":
                return await self._extract_pdf(file_path)
            elif file_ext == "docx":
                return await self._extract_docx(file_path)
            elif file_ext in ["txt", "md"]:
                return await self._extract_text_file(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    async def _extract_pdf(self, file_path: str) -> str:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    
    async def _extract_docx(self, file_path: str) -> str:
        from docx import Document
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    async def _extract_text_file(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
